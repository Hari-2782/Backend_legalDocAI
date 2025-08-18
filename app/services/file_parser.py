import hashlib
import os
import tempfile
import gc
from typing import List, Dict, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
import fitz  # PyMuPDF

class UniversalFileParser:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 100):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def calculate_file_hash(self, content: bytes) -> str:
        """Calculate SHA-256 hash of file content for duplicate detection"""
        return hashlib.sha256(content).hexdigest()
    
    def extract_text_from_file_bytes(self, file_bytes: bytes, filename: str, max_pages: int = None) -> Dict:
        """
        Extract text from supported file types: PDF, Word (.doc/.docx), and TXT files.
        Returns dict with text, hash, and metadata.
        """
        file_extension = os.path.splitext(filename)[1].lower()
        
        # Only support PDF, Word, and TXT files
        supported_extensions = ['.pdf', '.doc', '.docx', '.txt']
        if file_extension not in supported_extensions:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": f"Unsupported file type: {file_extension}. Only PDF, Word (.doc/.docx), and TXT files are supported.",
                "chunks": [],
                "total_chunks": 0,
                "file_type": file_extension
            }
        
        try:
            if file_extension == '.pdf':
                return self._extract_from_pdf(file_bytes, max_pages)
            elif file_extension == '.txt':
                return self._extract_from_text(file_bytes, filename)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_word(file_bytes, filename)
                
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": f"Error processing {file_extension} file: {str(e)}",
                "chunks": [],
                "total_chunks": 0,
                "file_type": file_extension
            }
    
    def _extract_from_pdf(self, pdf_bytes: bytes, max_pages: int = None) -> Dict:
        """Extract text from PDF using existing PDF parser logic"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_bytes)
                temp_file_path = temp_file.name
            
            try:
                loader = PyMuPDFLoader(temp_file_path)
                pages = loader.load()
                
                if max_pages is not None:
                    pages = pages[:max_pages]
                
                extracted_data = {
                    "hash": self.calculate_file_hash(pdf_bytes),
                    "pages": [],
                    "total_pages": len(pages),
                    "file_size": len(pdf_bytes),
                    "file_type": ".pdf"
                }
                
                for i, page in enumerate(pages):
                    page_data = {
                        "page": i + 1,
                        "text": page.page_content,
                        "metadata": page.metadata
                    }
                    extracted_data["pages"].append(page_data)
                    
                    if i % 10 == 0:
                        gc.collect()
                
                return extracted_data
                
            finally:
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(pdf_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": ".pdf"
            }
    
    def _extract_from_text(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError("Could not decode file with any supported encoding")
            
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "pages": [{"page": 1, "text": text, "metadata": {"filename": filename}}],
                "total_pages": 1,
                "file_size": len(file_bytes),
                "file_type": os.path.splitext(filename)[1].lower(),
                "full_text": text
            }
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
    
    def _extract_from_word(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from Word documents with optimized processing"""
        try:
            import docx
            from io import BytesIO
            
            # Optimized Word processing - extract text more efficiently
            doc = docx.Document(BytesIO(file_bytes))
            
            # Use list comprehension for faster text extraction
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            
            # Also extract text from tables if present (optional optimization)
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text.append(row_text)
            
            if tables_text:
                text += "\n\nTables:\n" + "\n".join(tables_text)
            
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "pages": [{"page": 1, "text": text, "metadata": {"filename": filename}}],
                "total_pages": 1,
                "file_size": len(file_bytes),
                "file_type": os.path.splitext(filename)[1].lower(),
                "full_text": text
            }
        except ImportError:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": "python-docx package required for Word document processing. Please install it: pip install python-docx",
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": f"Error processing Word document: {str(e)}",
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
    
    def _extract_from_html(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            html_content = file_bytes.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "pages": [{"page": 1, "text": text, "metadata": {"filename": filename}}],
                "total_pages": 1,
                "file_size": len(file_bytes),
                "file_type": os.path.splitext(filename)[1].lower(),
                "full_text": text
            }
        except ImportError:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": "beautifulsoup4 package required for HTML processing",
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
    
    def _extract_from_json(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from JSON files"""
        try:
            import json
            
            json_content = file_bytes.decode('utf-8')
            data = json.loads(json_content)
            text = json.dumps(data, indent=2)
            
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "pages": [{"page": 1, "text": text, "metadata": {"filename": filename}}],
                "total_pages": 1,
                "file_size": len(file_bytes),
                "file_type": os.path.splitext(filename)[1].lower(),
                "full_text": text
            }
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
    
    def _extract_from_csv(self, file_bytes: bytes, filename: str) -> Dict:
        """Extract text from CSV files"""
        try:
            import csv
            from io import StringIO
            
            csv_content = file_bytes.decode('utf-8')
            csv_reader = csv.reader(StringIO(csv_content))
            rows = list(csv_reader)
            
            # Convert CSV to readable text format
            text_lines = []
            for i, row in enumerate(rows):
                if i == 0:  # Header row
                    text_lines.append("Headers: " + " | ".join(row))
                else:
                    text_lines.append(f"Row {i}: " + " | ".join(row))
            
            text = "\n".join(text_lines)
            
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "pages": [{"page": 1, "text": text, "metadata": {"filename": filename}}],
                "total_pages": 1,
                "file_size": len(file_bytes),
                "file_type": os.path.splitext(filename)[1].lower(),
                "full_text": text
            }
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }
    
    def chunk_text_optimized(self, text: str) -> List[str]:
        """Use LangChain's optimized text splitting for better memory management."""
        try:
            if not text or len(text.strip()) < 50:
                return []
            
            chunks = self.text_splitter.split_text(text)
            filtered_chunks = [chunk.strip() for chunk in chunks if len(chunk.strip()) > 20]
            
            return filtered_chunks
            
        except Exception as e:
            print(f"Error chunking text: {e}")
            return []
    
    def process_file_in_batches(self, file_bytes: bytes, filename: str, batch_size: int = 3) -> Dict:
        """
        Process any file type in batches to minimize memory usage.
        """
        try:
            extracted_data = self.extract_text_from_file_bytes(file_bytes, filename)
            
            if "error" in extracted_data:
                return extracted_data
            
            all_chunks = []
            total_chunks = 0
            
            # Process pages in batches
            for i in range(0, len(extracted_data["pages"]), batch_size):
                batch_pages = extracted_data["pages"][i:i + batch_size]
                
                for page_data in batch_pages:
                    page_text = page_data["text"]
                    page_num = page_data["page"]
                    
                    chunks = self.chunk_text_optimized(page_text)
                    
                    for idx, chunk in enumerate(chunks):
                        chunk_data = {
                            "text": chunk,
                            "page": page_num,
                            "chunk_id": f"{extracted_data['hash']}::p{page_num}::c{idx}",
                            "metadata": {
                                "file_hash": extracted_data["hash"],
                                "page": page_num,
                                "chunk_index": idx,
                                "file_type": extracted_data.get("file_type", "unknown")
                            }
                        }
                        all_chunks.append(chunk_data)
                        total_chunks += 1
                
                gc.collect()
            
            extracted_data["chunks"] = all_chunks
            extracted_data["total_chunks"] = total_chunks
            
            return extracted_data
            
        except Exception as e:
            return {
                "hash": self.calculate_file_hash(file_bytes),
                "error": str(e),
                "chunks": [],
                "total_chunks": 0,
                "file_type": os.path.splitext(filename)[1].lower()
            }

# Create global instance
universal_parser = UniversalFileParser()

# Backward compatibility - keep existing pdf_parser for existing code
from .pdf_parser import pdf_parser
